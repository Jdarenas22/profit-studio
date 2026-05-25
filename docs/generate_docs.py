"""
Genera los archivos .docx de documentación de ProFit Studio.
Ejecutar desde la raíz del proyecto: python docs/generate_docs.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colores corporativos ────────────────────────────────────────────────────────
NARANJA    = RGBColor(0xFF, 0x6B, 0x00)   # #FF6B00
NEGRO      = RGBColor(0x1A, 0x1A, 0x1A)   # casi negro
GRIS_CLARO = RGBColor(0xF5, 0xF5, 0xF5)
GRIS_MED   = RGBColor(0x6B, 0x6B, 0x6B)
BLANCO     = RGBColor(0xFF, 0xFF, 0xFF)
VERDE      = RGBColor(0x16, 0xA3, 0x4A)


# ── Helpers ─────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Aplica color de fondo a una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Aplica bordes a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        val = kwargs.get(edge, {})
        if val:
            el = OxmlElement(f'w:{edge}')
            for k, v in val.items():
                el.set(qn(f'w:{k}'), v)
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, color=None, size=None, font='Calibri'):
    run = para.add_run(text)
    run.font.name = font
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def heading1(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run(text)
    run.font.name  = 'Calibri'
    run.font.bold  = True
    run.font.size  = Pt(18)
    run.font.color.rgb = NARANJA
    return para


def heading2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.font.name  = 'Calibri'
    run.font.bold  = True
    run.font.size  = Pt(13)
    run.font.color.rgb = NEGRO
    # Línea inferior naranja simulada con borde de párrafo
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'FF6B00')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def heading3(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(2)
    run = para.add_run(text)
    run.font.name  = 'Calibri'
    run.font.bold  = True
    run.font.size  = Pt(11)
    run.font.color.rgb = NARANJA
    return para


def body(doc, text, indent=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(4)
    if indent:
        para.paragraph_format.left_indent = Cm(0.8)
    run = para.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(10)
    run.font.color.rgb = NEGRO
    return para


def bullet(doc, text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent   = Cm(0.6 + level * 0.5)
    para.paragraph_format.space_before  = Pt(1)
    para.paragraph_format.space_after   = Pt(1)
    # Soporte para texto mixto: **negrita** y `code`
    _add_mixed_run(para, text)
    return para


def numbered(doc, text, level=0):
    para = doc.add_paragraph(style='List Number')
    para.paragraph_format.left_indent   = Cm(0.6 + level * 0.5)
    para.paragraph_format.space_before  = Pt(1)
    para.paragraph_format.space_after   = Pt(1)
    _add_mixed_run(para, text)
    return para


def _add_mixed_run(para, text):
    """Parsea **negrita**, `code` y texto normal dentro de un párrafo."""
    import re
    tokens = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            r = para.add_run(token[2:-2])
            r.font.bold = True
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
        elif token.startswith('`') and token.endswith('`'):
            r = para.add_run(token[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            r = para.add_run(token)
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            r.font.color.rgb = NEGRO


def code_block(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Cm(0.5)
    para.paragraph_format.right_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    # Fondo gris simulado con shading
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    run = para.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return para


def styled_table(doc, headers, rows, col_widths=None):
    """Tabla con cabecera naranja y filas alternadas."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Cabecera
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, 'FF6B00')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.bold  = True
        run.font.name  = 'Calibri'
        run.font.size  = Pt(9)
        run.font.color.rgb = BLANCO

    # Filas de datos
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'FDF3EC'  # blanco / naranja muy claro
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            _add_mixed_run(p, str(cell_text))
            for run in p.runs:
                run.font.size = Pt(9)

    # Anchos de columna
    if col_widths:
        for r in table.rows:
            for i, w in enumerate(col_widths):
                r.cells[i].width = Inches(w)

    doc.add_paragraph()  # espacio tras tabla
    return table


def add_cover(doc, title, subtitle, version, date):
    """Portada con fondo oscuro simulado."""
    # Espacio superior
    for _ in range(4):
        doc.add_paragraph()

    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ProFit Studio')
    run.font.name  = 'Calibri'
    run.font.bold  = True
    run.font.size  = Pt(32)
    run.font.color.rgb = NARANJA

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(title)
    r2.font.name  = 'Calibri'
    r2.font.bold  = True
    r2.font.size  = Pt(20)
    r2.font.color.rgb = NEGRO

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    r3.font.name   = 'Calibri'
    r3.font.italic = True
    r3.font.size   = Pt(12)
    r3.font.color.rgb = GRIS_MED

    for _ in range(2):
        doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(f'Versión {version}  ·  {date}')
    r4.font.name  = 'Calibri'
    r4.font.size  = Pt(10)
    r4.font.color.rgb = GRIS_MED

    doc.add_page_break()


def set_doc_margins(doc, top=2, bottom=2, left=2.5, right=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 1 — PRODUCT DOCUMENTATION (Agile / Product Owner)
# ═══════════════════════════════════════════════════════════════════════════════

def generate_product_doc():
    doc = Document()
    set_doc_margins(doc)

    add_cover(
        doc,
        title    = 'Documentación de Producto',
        subtitle = 'Metodología Ágil · Entrega para Product Owner',
        version  = '1.0',
        date     = 'Mayo 2026',
    )

    # ── 1. Visión del Producto ───────────────────────────────────────────────
    heading1(doc, '1. Visión del Producto')
    body(doc,
        'Para gimnasios y entrenadores personales que gestionan clientes de forma manual '
        '(WhatsApp, Excel), ProFit Studio es una plataforma web de gestión integral que '
        'centraliza clientes, membresías, rutinas, valoraciones físicas y pagos en un solo lugar.')
    body(doc,
        'A diferencia de herramientas genéricas como Google Sheets o apps de fitness masivas, '
        'nuestro producto está personalizado para el flujo de trabajo real de la entrenadora: '
        'desde el registro del cliente hasta su rutina diaria con videos de YouTube.')

    # ── 2. Usuarios del Sistema ──────────────────────────────────────────────
    heading1(doc, '2. Usuarios del Sistema (Roles)')
    styled_table(doc,
        headers=['Actor', 'Descripción', 'Nivel de acceso'],
        rows=[
            ['Yiseth (Superusuaria)', 'Dueña y entrenadora principal. Gestiona todo el sistema.',
             'Panel completo: clientes, pagos, entrenadores, configuración'],
            ['Entrenador/a', 'Staff contratado. Gestiona solo sus clientes asignados.',
             'Panel parcial: sin pagos, sin gestión de otros trainers'],
            ['Cliente (Miembro)', 'Usuario final del gimnasio.',
             'Dashboard personal: rutinas, valoraciones, pagos propios'],
            ['Visitante', 'Persona que llega desde redes sociales o búsqueda.',
             'Página pública: servicios, planes, contacto por WhatsApp'],
        ],
        col_widths=[1.4, 2.5, 2.7],
    )

    # ── 3. Épicas ────────────────────────────────────────────────────────────
    heading1(doc, '3. Épicas y Estado Actual')

    epicas = [
        {
            'num': 1, 'titulo': 'Presencia Web Pública', 'estado': 'Completada',
            'objetivo': 'Convertir visitantes en clientes potenciales.',
            'historias': [
                ('Como visitante, quiero ver los servicios del gimnasio para evaluar si me conviene.',
                 'Página home con hero, servicios, entrenadores y testimonios.', '✅'),
                ('Como visitante, quiero conocer al equipo de entrenadores y su experiencia.',
                 'Sección "Nosotros" con foto, nombre y bio de cada trainer.', '✅'),
                ('Como visitante, quiero ver los planes y precios para tomar una decisión.',
                 'Página de planes con precios en COP y botón de contratación.', '✅'),
                ('Como visitante, quiero contactar fácilmente al gimnasio.',
                 'Botón WhatsApp flotante en todas las páginas + página de contacto.', '✅'),
                ('Como visitante, quiero registrarme yo mismo sin necesidad de ir presencialmente.',
                 'Formulario de registro público con selección de plan de interés.', '✅'),
            ],
        },
        {
            'num': 2, 'titulo': 'Gestión de Clientes y Membresías', 'estado': 'Completada',
            'objetivo': 'Reemplazar el seguimiento en WhatsApp/Excel por una herramienta digital.',
            'historias': [
                ('Como entrenadora, quiero agregar un cliente nuevo desde el panel.',
                 'Formulario con datos personales y activación de membresía opcional.', '✅'),
                ('Como entrenadora, quiero activar o renovar la membresía de un cliente.',
                 'Vista de gestión con fecha inicio/fin calculada automáticamente.', '✅'),
                ('Como entrenadora, quiero saber qué membresías vencen en los próximos 7 días.',
                 'Dashboard con lista de membresías próximas a vencer.', '✅'),
                ('Como cliente, quiero ver el estado de mi membresía desde mi panel.',
                 'Dashboard del miembro con días restantes y estado de la membresía.', '✅'),
                ('Como cliente sin membresía vigente, quiero saber que mi acceso está limitado.',
                 'Página de aviso de membresía expirada con enlace a WhatsApp.', '✅'),
            ],
        },
        {
            'num': 3, 'titulo': 'Rutinas de Entrenamiento', 'estado': 'Completada',
            'objetivo': 'Dar a cada cliente una rutina personalizada accesible desde el celular.',
            'historias': [
                ('Como entrenadora, quiero crear rutinas con múltiples días y ejercicios.',
                 'Builder visual con días, series, reps y descanso por ejercicio.', '✅'),
                ('Como entrenadora, quiero asignar una rutina a un cliente específico.',
                 'Las rutinas se crean directamente para un usuario y aparecen en su panel.', '✅'),
                ('Como cliente, quiero ver mi rutina con la imagen y video de cada ejercicio.',
                 'Vista de rutina con thumbnail; clic en imagen abre YouTube.', '✅'),
                ('Como cliente, quiero registrar cuando completé un día de entrenamiento.',
                 'Botón "Completé este día" que guarda el registro de entrenamiento.', '✅'),
            ],
        },
        {
            'num': 4, 'titulo': 'Banco de Ejercicios', 'estado': 'Completada',
            'objetivo': 'Centralizar el contenido de entrenamiento reutilizable entre rutinas.',
            'historias': [
                ('Como entrenadora, quiero crear ejercicios con descripción, músculos y video.',
                 'Formulario con texto, imagen y URL de YouTube (no requiere subir video).', '✅'),
                ('Como entrenadora, quiero organizar ejercicios por categoría muscular.',
                 'Filtro por categoría en la lista (12 categorías precargadas).', '✅'),
                ('Como cliente, quiero ver el video de un ejercicio haciendo clic en su imagen.',
                 'Imagen con overlay de play que abre YouTube en pestaña nueva.', '✅'),
            ],
        },
        {
            'num': 5, 'titulo': 'Valoraciones Físicas y Mediciones', 'estado': 'Completada',
            'objetivo': 'Tener un historial científico del progreso de cada cliente.',
            'historias': [
                ('Como entrenadora, quiero registrar una valoración inicial con datos físicos.',
                 'Formulario con peso, talla, edad y objetivos. IMC calculado automáticamente.', '✅'),
                ('Como entrenadora, quiero aplicar el Test de Ruffier-Dickson.',
                 'Campos P0/P1/P2, índice IRD calculado con clasificación automática.', '✅'),
                ('Como entrenadora, quiero registrar mediciones periódicas de peso y medidas.',
                 'Historial de mediciones con peso, cintura e IMC en el perfil del cliente.', '✅'),
            ],
        },
        {
            'num': 6, 'titulo': 'Pagos', 'estado': 'Completada',
            'objetivo': 'Tener trazabilidad de todos los ingresos del gimnasio.',
            'historias': [
                ('Como cliente, quiero pagar mi membresía en línea con tarjeta, PSE o Nequi.',
                 'Checkout con Wompi; webhook activa membresía automáticamente al aprobar.', '✅'),
                ('Como entrenadora, quiero registrar pagos en efectivo o transferencia.',
                 'Formulario de pago manual con método, monto, fecha y comprobante.', '✅'),
                ('Como Yiseth, quiero ver el historial completo de pagos.',
                 'Lista de pagos manuales con filtro por cliente (solo visible a superusuaria).', '✅'),
            ],
        },
        {
            'num': 7, 'titulo': 'Panel de Entrenadores y Equipo', 'estado': 'Completada',
            'objetivo': 'Escalar el modelo de negocio con múltiples entrenadores.',
            'historias': [
                ('Como Yiseth, quiero crear cuentas para nuevos entrenadores.',
                 'Formulario de alta de trainer con datos y contraseña inicial.', '✅'),
                ('Como Yiseth, quiero asignar clientes a cada entrenador.',
                 'Vista de asignación con lista de trainers y conteo de clientes.', '✅'),
                ('Como Yiseth, quiero ver qué clientes no tienen entrenador asignado.',
                 'Alerta en dashboard y badge amarillo "Sin entrenador" en lista de clientes.', '✅'),
                ('Como Yiseth, quiero actualizar la foto y bio de cada entrenador para la web.',
                 'Edición de perfil desde panel; foto aparece en sección Nosotros.', '✅'),
                ('Como entrenador, quiero actualizar mi propio perfil y foto.',
                 'Botón "Mi perfil" en el nav del panel trainer.', '✅'),
            ],
        },
    ]

    for epica in epicas:
        heading2(doc, f"ÉPICA {epica['num']} — {epica['titulo']}  [{epica['estado']}]")
        body(doc, f"Objetivo: {epica['objetivo']}")
        styled_table(doc,
            headers=['Historia de Usuario', 'Criterio de Aceptación', 'Estado'],
            rows=epica['historias'],
            col_widths=[2.5, 2.8, 0.6],
        )

    # ── 4. Arquitectura Técnica ──────────────────────────────────────────────
    heading1(doc, '4. Arquitectura Técnica (Resumen Ejecutivo)')
    styled_table(doc,
        headers=['Capa', 'Tecnología'],
        rows=[
            ['Backend', 'Django 5.0.6, Python 3.14'],
            ['Base de datos', 'PostgreSQL (Railway)'],
            ['Almacenamiento media', 'Cloudflare R2 (S3-compatible)'],
            ['Frontend', 'Tailwind CSS + Alpine.js'],
            ['Pagos', 'Wompi (tarjeta, PSE, Nequi, Daviplata)'],
            ['Seguridad', 'django-axes (límite de intentos de login)'],
            ['Hosting / CD', 'Railway — deploy automático desde GitHub'],
        ],
        col_widths=[2.2, 4.0],
    )

    # ── 5. KPIs ──────────────────────────────────────────────────────────────
    heading1(doc, '5. Métricas del Producto (KPIs Sugeridos)')
    styled_table(doc,
        headers=['KPI', 'Cómo medirlo', 'Meta inicial'],
        rows=[
            ['Clientes activos', 'Membresías con estado Activa', 'Crecimiento mensual'],
            ['Tasa de renovación', 'Membresías renovadas / vencidas', '> 70 %'],
            ['Tiempo de respuesta al cliente nuevo', 'Desde registro hasta asignación de trainer', '< 24 horas'],
            ['Ejercicios con video', 'Ejercicios con URL de YouTube / total', '> 80 %'],
            ['Pagos digitales vs. manuales', 'Pagos Wompi aprobados / pagos manuales', 'Tendencia online'],
        ],
        col_widths=[2.0, 2.5, 1.8],
    )

    # ── 6. Backlog ───────────────────────────────────────────────────────────
    heading1(doc, '6. Backlog Priorizado (Próximas Iteraciones)')

    heading3(doc, 'Prioridad Alta')
    styled_table(doc,
        headers=['ID', 'Historia', 'Valor de negocio'],
        rows=[
            ['B-01', 'Como cliente, quiero recibir una notificación cuando mi membresía esté por vencer.',
             'Reduce pérdida de clientes por olvido de renovación.'],
            ['B-02', 'Como entrenadora, quiero exportar el listado de clientes a Excel/PDF.',
             'Reduce dependencia de herramientas externas para reportes.'],
            ['B-03', 'Como cliente, quiero ver mi progreso de peso en una gráfica.',
             'Aumenta engagement y percepción de valor del servicio.'],
        ],
        col_widths=[0.5, 3.2, 2.6],
    )

    heading3(doc, 'Prioridad Media')
    styled_table(doc,
        headers=['ID', 'Historia', 'Valor de negocio'],
        rows=[
            ['B-04', 'Como visitante, quiero pagar la membresía directamente desde la web.',
             'Reduce fricción en la conversión de leads.'],
            ['B-05', 'Como entrenadora, quiero enviar un mensaje grupal a clientes por vencer.',
             'Automatiza la gestión de retención.'],
            ['B-06', 'Como cliente, quiero calificar mis sesiones de entrenamiento.',
             'Feedback para la entrenadora sin reuniones.'],
        ],
        col_widths=[0.5, 3.2, 2.6],
    )

    heading3(doc, 'Deuda Técnica')
    styled_table(doc,
        headers=['ID', 'Descripción', 'Impacto'],
        rows=[
            ['T-01', 'Configurar Cloudflare R2 para almacenamiento permanente de imágenes.',
             'Las fotos se pierden en cada deploy sin R2.'],
            ['T-02', 'Agregar tests automatizados (pytest-django) al pipeline de CI.',
             'Previene regresiones al agregar nuevas funciones.'],
            ['T-03', 'Implementar paginación en listas de clientes y ejercicios.',
             'Necesario cuando superen ~100 registros.'],
        ],
        col_widths=[0.5, 3.2, 2.6],
    )

    # ── 7. Definition of Done ────────────────────────────────────────────────
    heading1(doc, '7. Definition of Done (DoD)')
    body(doc, 'Una historia se considera terminada cuando:')
    criterios = [
        'La funcionalidad está implementada y funciona en producción (Railway)',
        'Los datos del negocio (nombres, precios, contacto) son reales, no placeholders',
        'El diseño es coherente con el sistema visual (dark theme, color naranja #FF6B00)',
        'El acceso está protegido según el rol del usuario (visitante / miembro / trainer / superusuaria)',
        'Los datos sensibles (contraseñas, claves API) están en variables de entorno, no en el código',
        'El cambio está en el repositorio GitHub y desplegado automáticamente',
    ]
    for c in criterios:
        bullet(doc, f'☐  {c}')

    # ── 8. Riesgos ───────────────────────────────────────────────────────────
    heading1(doc, '8. Riesgos Identificados')
    styled_table(doc,
        headers=['Riesgo', 'Probabilidad', 'Impacto', 'Mitigación'],
        rows=[
            ['Imágenes se pierden sin R2 configurado', 'Alta', 'Medio',
             'Configurar Cloudflare R2 (guía disponible)'],
            ['Deploy fallido deja la app sin trainer superusuario', 'Baja', 'Alto',
             'create_trainer en start.sh es idempotente'],
            ['Credenciales Wompi expuestas en el código', 'Baja', 'Crítico',
             'Variables de entorno en Railway; nunca en código'],
            ['Crecimiento de clientes satura la base de datos', 'N/A', 'N/A',
             'Ya usa PostgreSQL desde el inicio'],
        ],
        col_widths=[2.0, 1.1, 0.9, 2.3],
    )

    # ── 9. Contacto ──────────────────────────────────────────────────────────
    heading1(doc, '9. Contacto del Equipo')
    styled_table(doc,
        headers=['Rol', 'Responsabilidad'],
        rows=[
            ['Product Owner', 'Yiseth Misas García (ProFit Studio)'],
            ['Desarrollo', 'Claude Code + ProFit Studio Team'],
            ['Repositorio', 'github.com/Jdarenas22/profit-studio'],
            ['Producción', 'Railway — deploy automático desde rama main'],
        ],
        col_widths=[2.0, 4.3],
    )

    path = 'docs/ProFitStudio_Documentacion_Agile_PO.docx'
    doc.save(path)
    print(f'[OK] Generado: {path}')


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 2 — CLAUDE.md técnico (guía para desarrolladores)
# ═══════════════════════════════════════════════════════════════════════════════

def generate_technical_doc():
    doc = Document()
    set_doc_margins(doc)

    add_cover(
        doc,
        title    = 'Guía Técnica del Repositorio',
        subtitle = 'Referencia para desarrolladores · Claude Code',
        version  = '1.0',
        date     = 'Mayo 2026',
    )

    # ── Intro ────────────────────────────────────────────────────────────────
    heading1(doc, 'Proyecto')
    body(doc,
        'ProFit Studio es una plataforma web de gestión para gimnasio/entrenamiento personal. '
        'Stack: Django 5.0.6 desplegado en Railway con PostgreSQL. El negocio opera en Colombia '
        '(moneda COP, zona horaria America/Bogota).')

    # ── Comandos ─────────────────────────────────────────────────────────────
    heading1(doc, 'Comandos Esenciales')

    heading2(doc, 'Desarrollo local')
    code_block(doc,
        'python manage.py runserver\n'
        'python manage.py migrate\n'
        'python manage.py load_initial_data          # categorías + planes de membresía\n'
        'python manage.py create_trainer \\           # crear Yiseth como superusuaria\n'
        '    --username yiseth \\                     # idempotente (no sobreescribe si existe)\n'
        '    --first-name Yiseth \\\n'
        '    --last-name "Misas García" \\\n'
        '    --email profitstudio075@gmail.com \\\n'
        '    --password "xxx" --superuser')

    heading2(doc, 'Producción (Railway)')
    body(doc,
        'El archivo start.sh ejecuta automáticamente: migrate → collectstatic → '
        'create_trainer → load_initial_data → gunicorn. Railway lo llama en cada deploy.')

    # ── Arquitectura de apps ─────────────────────────────────────────────────
    heading1(doc, 'Arquitectura')

    heading2(doc, 'Estructura de apps (/apps/)')
    styled_table(doc,
        headers=['App', 'Responsabilidad'],
        rows=[
            ['accounts', 'Usuario custom, auth, dashboard, panel de entrenadores'],
            ['memberships', 'Planes y membresías (activar, renovar, vencer)'],
            ['exercises', 'Banco de ejercicios con imagen/video YouTube'],
            ['routines', 'Rutinas multi-día asignadas a clientes'],
            ['assessments', 'Valoración inicial (IMC + Test Ruffier-Dickson) y mediciones corporales'],
            ['payments', 'Pagos online (Wompi) y manuales (efectivo/transferencia)'],
            ['public', 'Páginas públicas (home, nosotros, contacto, testimonios)'],
        ],
        col_widths=[1.5, 4.8],
    )

    # ── Modelo de roles ──────────────────────────────────────────────────────
    heading2(doc, 'Modelo de Roles')
    body(doc,
        'El modelo User extiende AbstractUser con campo role ("trainer" / "member"). '
        'Dentro de los trainers, is_superuser=True identifica a Yiseth (dueña). '
        'Esta distinción determina:')
    bullet(doc, 'Solo **is_superuser** ve Pagos y Entrenadores en el nav del panel trainer')
    bullet(doc, 'Solo **is_superuser** puede crear/editar otros trainers y asignar clientes')
    bullet(doc, '**is_superuser** ve todos los clientes; otros trainers solo ven los asignados')
    doc.add_paragraph()
    body(doc, 'Propiedades clave del modelo User:')
    code_block(doc,
        'user.is_trainer             # True si role == "trainer"\n'
        'user.has_active_membership  # True si membership.is_valid\n'
        'user.assigned_trainer       # FK a otro User (trainer asignado)\n'
        'user.assigned_clients       # reverse relation (clientes de este trainer)')

    # ── Decoradores ──────────────────────────────────────────────────────────
    heading2(doc, 'Decoradores de Acceso')
    code_block(doc,
        '@trainer_required        # is_authenticated + is_trainer; 403 si no\n'
        '@membership_required     # trainers pasan siempre; members necesitan membresía válida\n'
        '@login_required          # estándar Django')
    body(doc,
        'Para vistas solo de superusuario no hay decorador dedicado — se usa '
        '"if not request.user.is_superuser: redirect(\'trainer_dashboard\')" al inicio de la vista.')

    # ── Settings ─────────────────────────────────────────────────────────────
    heading2(doc, 'Settings')
    code_block(doc,
        'config/settings/\n'
        '├── base.py         # Base común (no define DEBUG ni ALLOWED_HOSTS)\n'
        '└── production.py   # Extiende base; fuerza DEBUG=False, configura R2, PostgreSQL')
    body(doc,
        'No existe local.py. La selección del settings file es via DJANGO_SETTINGS_MODULE '
        '(default en start.sh: config.settings.production).')

    heading3(doc, 'Variables de entorno requeridas en Railway')
    styled_table(doc,
        headers=['Variable', 'Descripción'],
        rows=[
            ['SECRET_KEY, DATABASE_URL, ALLOWED_HOSTS', 'Configuración base de Django'],
            ['TRAINER_PASSWORD', 'Contraseña inicial de Yiseth (superusuaria)'],
            ['R2_ACCESS_KEY_ID, R2_SECRET_ACCESS_KEY', 'Credenciales API Cloudflare R2'],
            ['R2_BUCKET_NAME, R2_ENDPOINT_URL, R2_PUBLIC_URL', 'Configuración del bucket R2'],
            ['WOMPI_PUBLIC_KEY, WOMPI_PRIVATE_KEY', 'Claves Wompi para pagos online'],
            ['WOMPI_INTEGRITY_SECRET, WOMPI_EVENTS_SECRET', 'Seguridad webhook Wompi'],
            ['EMAIL_HOST_USER, EMAIL_HOST_PASSWORD', 'SMTP para recuperación de contraseña'],
            ['WHATSAPP_NUMBER, WHATSAPP_LINK, INSTAGRAM_URL', 'Datos de contacto del negocio'],
        ],
        col_widths=[3.0, 3.3],
    )

    # ── Almacenamiento ───────────────────────────────────────────────────────
    heading2(doc, 'Almacenamiento de Archivos Media')
    bullet(doc, '**Sin R2 configurado:** guarda en BASE_DIR/media/ (se pierde con cada redeploy en Railway)')
    bullet(doc, '**Con R2:** usa django-storages + boto3 con AWS_QUERYSTRING_AUTH=False (URLs permanentes)')
    body(doc,
        'R2_ENDPOINT_URL es el endpoint S3 API de boto3 para subir archivos. '
        'R2_PUBLIC_URL es la URL pública del bucket (ej: https://pub-xxx.r2.dev) '
        'y se usa como MEDIA_URL para servir las imágenes en la web.')

    # ── Frontend ─────────────────────────────────────────────────────────────
    heading2(doc, 'Frontend')
    styled_table(doc,
        headers=['Elemento', 'Detalle'],
        rows=[
            ['CSS', 'Tailwind CDN — sin build step'],
            ['JS interactivo', 'Alpine.js (x-data, x-model, x-show)'],
            ['Tema', 'Dark — bg-zinc-900/bg-black — acento #FF6B00 (text-primary, bg-primary)'],
            ['Base templates', 'base.html (público/miembros) y trainer/base.html (panel admin)'],
            ['Context processor global', 'Inyecta WHATSAPP_LINK, WHATSAPP_NUMBER, INSTAGRAM_URL en todos los templates'],
        ],
        col_widths=[2.0, 4.3],
    )

    # ── Pagos ────────────────────────────────────────────────────────────────
    heading2(doc, 'Pagos (Dos Sistemas Independientes)')
    numbered(doc,
        '**Wompi** (apps/payments): pagos online con webhook. '
        'La vista wompi_webhook es @csrf_exempt. '
        'Payment se crea al iniciar checkout; el webhook actualiza el status y activa la membresía.')
    numbered(doc,
        '**ManualPayment** (apps/payments): registrado por trainers para pagos en efectivo/transferencia. '
        'Solo visible a is_superuser en el nav.')

    # ── Cálculos automáticos ─────────────────────────────────────────────────
    heading2(doc, 'Cálculos Automáticos en save()')
    bullet(doc, '`InitialAssessment.save()` → calcula y guarda imc e imc_classification')
    bullet(doc, '`DixonTest.save()` → calcula index_value y classification (IRD = (P0+P1+P2-200)/10)')
    bullet(doc, '`BodyMeasurement.save()` → calcula imc e imc_classification')

    # ── Media de ejercicios ──────────────────────────────────────────────────
    heading2(doc, 'Ejercicios: Prioridad de Media')
    body(doc, 'El template exercises/detail.html evalúa en este orden:')
    numbered(doc, '`image` + `video_url` → imagen clickeable con overlay play que abre YouTube (tab nueva)')
    numbered(doc, 'Solo `video_url` → iframe embed de YouTube')
    numbered(doc, 'Solo `video_file` → player HTML5')
    numbered(doc, 'Solo `image` → imagen estática')
    numbered(doc, 'Nada → placeholder con icono de mancuerna')
    body(doc,
        'El modelo Exercise tiene una property youtube_embed_url que convierte cualquier '
        'URL de YouTube (youtu.be, watch?v=, shorts/) al formato de embed.')

    # ── Rutinas ──────────────────────────────────────────────────────────────
    heading2(doc, 'Rutinas')
    body(doc,
        'Estructura jerárquica: Routine → RoutineDay (días) → RoutineExercise '
        '(ejercicios con sets/reps/descanso/observaciones). '
        'El builder del trainer es interactivo (Alpine.js). '
        'Los clientes ven sus rutinas en member_routine.html con thumbnails de imagen '
        'y botón de video YouTube directo.')

    path = 'docs/ProFitStudio_Guia_Tecnica_Desarrolladores.docx'
    doc.save(path)
    print(f'[OK] Generado: {path}')


# ── Main ────────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    generate_product_doc()
    generate_technical_doc()
    print('\n[LISTO] Ambos documentos generados en la carpeta docs/')
